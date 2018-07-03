import os
from docx import Document

document = Document()
ROOT_DIR = os.path.dirname(os.path.abspath(__file__))
print(ROOT_DIR)

path = 'C:\\Users\\yespindola\\Documents\\MonthlyTaskReportGeneration\\sources\\'

# path_to_file = "Report - Sept - 2017.txt"
# performance_file = open(path_to_file, 'w')

for filename in os.listdir(path):
    print(filename)
    fields = filename.split(" - ")
    print(filename)
    date = fields[0]
    name = fields[1]
    # print(date)
    # print(name)
    date_fields = date.split("-")
    date = date_fields[0]
    date_fields_aux = date.split("_")
    year = date_fields_aux[0]
    month = date_fields_aux[1]
    if month == "07":
        month = "julio"
    elif month == "08":
        month = "agosto"
    elif month == "09":
        month = "septiembre"
    elif month == "10":
        month = "octubre"
    else:
        month = "no determinado"
    day = date_fields_aux[2]
    print(name)
    if "Performance" in name or "food" in name or "words" in name or "server" in name:
        print("Ejecución")
        name_fields = name.split("_")
        if len(name_fields) == 2:
            component_name = name_fields[0]
        else:
            component_name = name_fields[0] + "_" + name_fields[1]
        if component_name == "CL":
            component_name = "ConversationLogic"
        elif component_name == "Analytics_L+KMeans":
            component_name = "Clustering (Louvain y KMeans)"
        elif component_name == "UL_A":
            component_name = "UserLogic"
        elif component_name == "FoodDetection_Server":
            component_name = "FoodDetection"
        elif component_name == "SeedListGenerator_Server" or component_name == "SeedListGenerator":
            component_name = "SeedListGenerator en FoodDetection"
        else:
            component_name = component_name
        # print(component_name)
        time = date_fields[1]
        time_fields = time.split("_")
        # print(time_fields)
        hour = time_fields[0]
        minute = time_fields[1]
        second = time_fields[2]
        if "Server" in name or "server" in name:
            location = "en el servidor Hércules"
        else:
            location = "de forma local"
        if "Performance" in name:
            register_type = "execution"
        else:
            register_type = "result"
        if register_type == "execution":
            description = "Resultados de rendimiento de la prueba " + location + " para el componente " \
                          + component_name + ", de forma automática ejecutada el " + day + " de " + month + " de " + year \
                          + " a las " + hour + ":" + minute + ":" + second + "."
        else:
            if "anagrams" in name:
                result_name = "Anagramas con palabras de comida"
            elif "hashtags_about_food" in name:
                result_name = "Hashtags relacionados con comida"
            elif "hashtags_with_food" in name:
                result_name = "Hashtags con palabras de comida"
            elif "user_mentions_about_food" in name:
                result_name = "Alias de usuarios relacionados con comida"
            elif "user_mentions_with_food" in name:
                result_name = "Alias de usuarios con palabras de comida"
            elif "what_words" in name:
                result_name = "Palabras de comida identificadas"
            elif "text_about_food" in name:
                result_name = "Texto con palabras de comida"
            elif "text_not_about_food" in name:
                result_name = "Texto sin palabras de comida"
            else:
                result_name = "Archivo no reconocido"
            description = result_name + " resultado de la ejecución " + location \
                          + " por medio del componente FoodDetection generado el " + day \
                          + " de " + month + " de " + year + " a las " + hour + ":" + minute + ":" + second + "."

    else:
        name_fields = name.split(".")
        doc_type = name_fields[len(name_fields)-1]
        print(doc_type)
        print(name)
        if doc_type == "PNG" or doc_type == "png":
            description = name + ". Versión del " + day + " de " + month + " de " + year + "."
        elif doc_type == "py":
            description = "Desarrollo: " + name + ". Versión del " + day + " de " + month + " de " + year + "."
        else:
            description = "Documento: " + name + ". Versión del " + day + " de " + month + " de " + year + "."
    print(description)
    p = document.add_paragraph()
    run = p.add_run(filename+ ": ")
    run.bold = True
    run.underline = True
    p.add_run(description)

document.save('Report - Oct - 2017.docx')
