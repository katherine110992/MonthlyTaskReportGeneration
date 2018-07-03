import os
from docx import Document

document = Document()
ROOT_DIR = os.path.dirname(os.path.abspath(__file__))
print(ROOT_DIR)

path = 'C:\\Users\\yespindola\\Documents\\MonthlyTaskReportGeneration\\sources\\'


for filename in os.listdir(path):
    print(filename)
    fields = filename.split(" - ")
    # print(fields)
    date = fields[0]
    name = fields[1]
    print(fields)
    print(date)
    print(name)
    date_fields = date.split("-")
    print(date_fields)
    date = date_fields[0]
    date_fields_aux = date.split("_")
    year = date_fields_aux[0]
    month = date_fields_aux[1]
    if month == "11":
        month = "noviembre"
    elif month == "12":
        month = "diciembre"
    elif month == "01":
        month = "enero"
    elif month == "02":
        month = "febrero"
    elif month == "03":
        month = "marzo"
    else:
        month = "no determinado"
    day = date_fields_aux[2]
    print(name)
    # falta if server o local entonces algo pasa para esos archivos que son resultados ya sea en el servidor o locales
    if "Server" in name or "server" in name:
        location = "en el servidor Hércules"
    else:
        location = "de forma local"
    if "Performance" in name:
        print('execution')
        # It is a execution
        name_fields = name.split("_")
        if len(name_fields) == 2:
            component_name = name_fields[0]
        else:
            component_name = name_fields[0] + "_" + name_fields[1]
        if component_name == "CL":
            component_name = "ConversationLogic"
        elif component_name == "Analytics_L+KMeans":
            component_name = "Clustering (Louvain y KMeans)"
        elif component_name == "UL_A":
            component_name = "UserLogic"
        elif component_name == "FoodDetection_Server":
            component_name = "FoodDetection"
        elif "DetectRegexSpecialEntitiesRawData" in component_name:
            component_name = "de detección de special_entities con expresiones regulares con datos en bruto en FoodDetection"
        elif "DetectEmoticonsWithRegex" in component_name:
            component_name = "de detección de emoticones con expresiones regulares en FoodDetection"
        elif "SpecialEntityTextAnalysisTest" in component_name:
            component_name = "de prueba de funcionalidad special_entity_detection de TextAnalysis desde ConversationLogic"
        elif "SyntaxPreparationTextAnalysisTest" in component_name:
            component_name = "de prueba de funcionalidad syntax_preparation de TextAnalysis desde ConversationLogic"
        elif "TextAnalysis" in component_name:
            component_name = "TextAnalysis"
        elif "SpacyTagMap" in component_name:
            component_name = "de pruebas de la librería spaCy en TextAnalysis"
        elif "Validation" in component_name:
            component_name = "de validación de emojis y emoticones detectados con TextAnalysis desde ConversationLogic"
        else:
            component_name = component_name
        # print(component_name)
        time = date_fields[1]
        time_fields = time.split("_")
        # print(time_fields)
        hour = time_fields[0]
        minute = time_fields[1]
        second = time_fields[2]
        description = "Resultados de rendimiento de la prueba " + location + " para el componente " \
                      + component_name + ", de forma automática ejecutada el " + day + " de " + month + " de " + year \
                      + " a las " + hour + ":" + minute + ":" + second + "."
    elif "list sources" in name:
        description = "Fuentes de información para la generación de listas. Versión del " + day \
                      + " de " + month + " de " + year + "."
    elif "list" in name:
        list_name = fields[2]
        if "stemmed_what_food" in list_name:
            result_name = "lista de comida stemizada"
        elif "tagged_what_food" in list_name:
            result_name = "lista de comida etiquetada con POS"
        elif "original_stemmed_what_food" in list_name:
            result_name = "lista de comida original y stemizada"
        elif "what_food" in list_name:
            result_name = "lista de comida mejorada"
        elif "raw_emojis" in list_name:
            result_name = "lista de emojis original conseguida de http://www.unicode.org/emoji/charts/full-emoji-list.html"
        elif "raw_emoticons" in list_name:
            result_name = "lista de emoticones original conseguida de http://cool-smileys.com/text-emoticons"
        elif "unicode_emojis" in list_name:
            result_name = "lista de emojis generada de la lista original"
        elif "emoticons" in list_name:
            result_name = "lista de emoticones generada de la lista original"
        elif "raw_complementary_characters" in list_name:
            result_name = "lista de caracteres complementarios conseguida de " \
                          "http://www.fileformat.info/info/unicode/category/So/list.htm y " \
                          "https://codepoints.net/basic_multilingual_plane"
        elif "complementary_characters" in list_name:
            result_name = "lista de caracteres complementarios generada de la lista original"
        elif "emoji_sentiment_data" in list_name:
            result_name = "lista de emojis con polaridad generada de fuente externa"
        elif "emoji_sentiment" in list_name:
            result_name = "lista de emojis con polaridad generada de fuente externa"
        elif "food_hashtag" in list_name:
            result_name = "lista de hashtags de alimentación para el componente TextAnalysis"
        elif "food_entity" in list_name:
            result_name = "lista de entidades de alimentación para el componente TextAnalysis"
        elif "food_usermention" in list_name:
            result_name = "lista de menciones de alimentación para el componente TextAnalysis"
        elif "food" in list_name:
            result_name = "lista de n-gramas de alimentación para el componente TextAnalysis"
        else:
            result_name = "Archivo no reconocido"
        description = "Última versión de la " + result_name + " del " + day \
                      + " de " + month + " de " + year + "."
    elif "Installation" in name or "installation" in name or 'instalación' in name:
        print('instalación')
        if 'BeautifulSoap' in name:
            library = 'BeautifulSoup'
        elif 'Spacy' in name:
            library = 'Spacy'
        else:
            library = "Libreria no reconocida"
        description = "Instalación de la librería " + library + " "+ location + " el " + day + " de " \
                      + month + " de " + year
    elif "users" in name or "conversations" in name:
        print('DAO results')
        if "conversations_per_days" in name:
            result_name = "Número total de conversaciones por día de la semana"
        elif "conversations_per_hour" in name:
            result_name = "Número total de conversaciones por hora del día"
        elif "users" in name:
            result_name = "Total de usuarios por mes"
        else:
            result_name = "Archivo no reconocido"
        # print(component_name)
        time = date_fields[1]
        time_fields = time.split("_")
        # print(time_fields)
        hour = time_fields[0]
        minute = time_fields[1]
        second = time_fields[2]
        description = result_name + " resultado de la ejecución " + location \
                      + " del componente DataAccessObject para la actualización de datos en la presentación a" \
                        " Grupo Nutresa, generado el " + day + " de " + month + " de " + year + " a las " \
                      + hour + ":" + minute + ":" + second + "."
    else:
        print('Docs')
        name_fields = name.split(".")
        doc_type = name_fields[len(name_fields)-1]
        print(doc_type)
        print(name)
        if doc_type == "PNG" or doc_type == "png":
            description = name + ". Versión del " + day + " de " + month + " de " + year + "."
        elif doc_type == "py":
            description = "Desarrollo: " + name + ". Versión del " + day + " de " + month + " de " + year + "."
        else:
            description = "Documento: " + name + ". Versión del " + day + " de " + month + " de " + year + "."
    print(description)
    p = document.add_paragraph()
    run = p.add_run(filename+ ": ")
    run.bold = True
    run.underline = True
    p.add_run(description)

document.save('Report - Feb - 2018.docx')




