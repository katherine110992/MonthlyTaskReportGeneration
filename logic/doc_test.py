import docx
from docx import Document

document = Document()

paragraph = document.add_paragraph('Lorem ipsum dolor sit amet.')
run = paragraph.add_run('dolor')
run.bold = True
run.italic = True
paragraph.add_run(' sit amet.')

p = document.add_paragraph()
p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True

p = document.add_paragraph()
p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True
run.bold = True
run.underline = True
document.save('demo.docx')