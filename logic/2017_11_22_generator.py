import os
from docx import Document

document = Document()
ROOT_DIR = os.path.dirname(os.path.abspath(__file__))
print(ROOT_DIR)

path = 'C:\\Users\\yespindola\\Documents\\MonthlyTaskReportGeneration\\sources\\'

# path_to_file = "Report - Sept - 2017.txt"
# performance_file = open(path_to_file, 'w')

for filename in os.listdir(path):
    print(filename)
    fields = filename.split(" - ")
    # print(fields)
    date = fields[0]
    name = fields[1]
    print(date)
    print(name)
    date_fields = date.split("-")
    print(date_fields)
    date = date_fields[0]
    date_fields_aux = date.split("_")
    year = date_fields_aux[0]
    month = date_fields_aux[1]
    if month == "11":
        month = "noviembre"
    elif month == "08":
        month = "agosto"
    elif month == "09":
        month = "septiembre"
    elif month == "10":
        month = "octubre"
    else:
        month = "no determinado"
    day = date_fields_aux[2]
    print(name)
    # falta if server o local entonces algo pasa para esos archivos que son resultados ya sea en el servidor o locales
    if "Server" in name or "server" in name:
        location = "en el servidor Hércules"
    else:
        location = "de forma local"
    if "Performance" in name:
        print('execution')
        # It is a execution
        name_fields = name.split("_")
        if len(name_fields) == 2:
            component_name = name_fields[0]
        else:
            component_name = name_fields[0] + "_" + name_fields[1]
        if component_name == "CL":
            component_name = "ConversationLogic"
        elif component_name == "Analytics_L+KMeans":
            component_name = "Clustering (Louvain y KMeans)"
        elif component_name == "UL_A":
            component_name = "UserLogic"
        elif component_name == "FoodDetection_Server":
            component_name = "FoodDetection"
        elif component_name == "SeedListGenerator_Server" or component_name == "SeedListGenerator":
            component_name = "SeedListGenerator en FoodDetection"
        elif component_name == "TaggedFoodListGeneration_Server" or component_name == "TaggedFoodListGeneration":
            component_name = "TaggedFoodListGeneration en FoodDetection"
        elif component_name == "StemmedFoodListGeneration_Server" or component_name == "StemmedFoodListGeneration":
            component_name = "StemmedFoodListGeneration en FoodDetection"
        elif component_name == "SpacyTagMap_Server" or component_name == "SpacyTagMap":
            component_name = "SpacyTagMap en FoodDetection"
        elif component_name == "Original&StemmedListGeneration_Server" or component_name == "Original&StemmedListGeneration":
            component_name = "Original&StemmedListGeneration en FoodDetection"
        elif component_name == "DetectFoodWithSpacy_Server" or component_name == "DetectFoodWithSpacy":
            component_name = "DetectFoodWithSpacy en FoodDetection"
        elif component_name == "EmojisUnicodeListGeneration_Server" or component_name == "EmojisUnicodeListGeneration":
            component_name = "EmojisUnicodeListGeneration en FoodDetection"
        elif component_name == "EmoticonListGeneration_Server" or component_name == "EmoticonListGeneration":
            component_name = "EmoticonListGeneration en FoodDetection"
        elif component_name == "DetectEmoticonsAndEmojis_Server" or component_name == "DetectEmoticonsAndEmojis":
            component_name = "DetectEmoticonsAndEmojis en FoodDetection"
        elif component_name == "AssignationEmojisListClassification_Server" or component_name == "AssignationEmojisListClassification":
            component_name = "AssignationEmojisListClassification en FoodDetection"
        elif component_name == "DetectLanguage_Server" or component_name == "DetectLanguage":
            component_name = "DetectLanguage en FoodDetection"
        else:
            component_name = component_name
        # print(component_name)
        time = date_fields[1]
        time_fields = time.split("_")
        # print(time_fields)
        hour = time_fields[0]
        minute = time_fields[1]
        second = time_fields[2]
        description = "Resultados de rendimiento de la prueba " + location + " para el componente " \
                      + component_name + ", de forma automática ejecutada el " + day + " de " + month + " de " + year \
                      + " a las " + hour + ":" + minute + ":" + second + "."
    elif "food" in name or "words" in name:
        print('FoodDetection results')
        if "anagrams" in name:
            result_name = "Anagramas con palabras de comida"
        elif "hashtags_about_food" in name:
            result_name = "Hashtags relacionados con comida"
        elif "hashtags_with_food" in name:
            result_name = "Hashtags con palabras de comida"
        elif "user_mentions_about_food" in name:
            result_name = "Alias de usuarios relacionados con comida"
        elif "user_mentions_with_food" in name:
            result_name = "Alias de usuarios con palabras de comida"
        elif "what_words" in name:
            result_name = "Palabras de comida identificadas"
        elif "new_what_words" in name:
            result_name = "Nuevas palabras de comida potenciales"
        elif "text_about_food" in name:
            result_name = "Texto con palabras de comida"
        elif "text_not_about_food" in name:
            result_name = "Texto sin palabras de comida"
        else:
            result_name = "Archivo no reconocido"
        # print(component_name)
        time = date_fields[1]
        time_fields = time.split("_")
        # print(time_fields)
        hour = time_fields[0]
        minute = time_fields[1]
        second = time_fields[2]
        description = result_name + " resultado de la ejecución " + location \
                      + " del componente FoodDetection generado el " + day \
                      + " de " + month + " de " + year + " a las " + hour + ":" + minute + ":" + second + "."
    elif "list sources" in name:
        description = "Fuentes de información para la generación de listas. Versión del " + day \
                      + " de " + month + " de " + year + "."
    elif "list" in name:
        print('lista')
        list_name = fields[2]
        if "stemmed_what_food" in list_name:
            result_name = "lista de comida stemizada"
        elif "tagged_what_food" in list_name:
            result_name = "lista de comida etiquetada con POS"
        elif "original_stemmed_what_food" in list_name:
            result_name = "lista de comida original y stemizada"
        elif "what_food" in list_name:
            result_name = "lista de comida mejorada"
        elif "raw_emojis" in list_name:
            result_name = "lista de emojis original conseguida de http://www.unicode.org/emoji/charts/full-emoji-list.html"
        elif "raw_emoticons" in list_name:
            result_name = "lista de emoticones original conseguida de http://cool-smileys.com/text-emoticons"
        elif "unicode_emojis" in list_name:
            result_name = "lista de emojis generada de la lista original"
        elif "emoticons" in list_name:
            result_name = "lista de emoticones generada de la lista original"
        elif "raw_complementary_characters" in list_name:
            result_name = "lista de caracteres complementarios conseguida de " \
                          "http://www.fileformat.info/info/unicode/category/So/list.htm y " \
                          "https://codepoints.net/basic_multilingual_plane"
        elif "complementary_characters" in list_name:
            result_name = "lista de caracteres complementarios generada de la lista original"
        else:
            result_name = "Archivo no reconocido"
        description = "Última versión de la " + result_name + " del " + day \
                      + " de " + month + " de " + year + "."
    elif "Installation" in name or "installation" in name or 'instalación' in name:
        print('instalación')
        if 'BeautifulSoap' in name:
            library = 'BeautifulSoup'
        elif 'Spacy' in name:
            library = 'Spacy'
        else:
            library = "Libreria no reconocida"
        description = "Instalación de la librería " + library + " "+ location + " el " + day + " de " \
                      + month + " de " + year
    elif "users" in name or "conversations" in name:
        print('DAO results')
        if "conversations_per_days" in name:
            result_name = "Número total de conversaciones por día de la semana"
        elif "conversations_per_hour" in name:
            result_name = "Número total de conversaciones por hora del día"
        elif "users" in name:
            result_name = "Total de usuarios por mes"
        else:
            result_name = "Archivo no reconocido"
        # print(component_name)
        time = date_fields[1]
        time_fields = time.split("_")
        # print(time_fields)
        hour = time_fields[0]
        minute = time_fields[1]
        second = time_fields[2]
        description = result_name + " resultado de la ejecución " + location \
                      + " del componente DataAccessObject para la actualización de datos en la presentación a" \
                        " Grupo Nutresa, generado el " + day + " de " + month + " de " + year + " a las " \
                      + hour + ":" + minute + ":" + second + "."
    else:
        print('Docs')
        name_fields = name.split(".")
        doc_type = name_fields[len(name_fields)-1]
        print(doc_type)
        print(name)
        if doc_type == "PNG" or doc_type == "png":
            description = name + ". Versión del " + day + " de " + month + " de " + year + "."
        elif doc_type == "py":
            description = "Desarrollo: " + name + ". Versión del " + day + " de " + month + " de " + year + "."
        else:
            description = "Documento: " + name + ". Versión del " + day + " de " + month + " de " + year + "."
    print(description)
    p = document.add_paragraph()
    run = p.add_run(filename+ ": ")
    run.bold = True
    run.underline = True
    p.add_run(description)

document.save('Report - Nov - 2017.docx')




