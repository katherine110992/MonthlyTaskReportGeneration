import os
from docx import Document

document = Document()
ROOT_DIR = os.path.dirname(os.path.abspath(__file__))
print(ROOT_DIR)

path = 'C:\\Users\\yespindola\\Documents\\MonthlyTaskReportGeneration\\sources\\'

counter = {
    'en el servidor Hércules': {
        'ConversationLogic': 0,
        'Clustering (Louvain y KMeans)': 0,
        'UserLogic': 0,
        'DataAccessObject': 0,
        'TextAnalysis': 0,
        'WordFrequencyAnalysis': 0,
        'StatisticalAnalysis': 0,
        'FileGeneration': 0,
        'ScheduledTaskManagement': 0
    },
    'de forma local':{
        'ConversationLogic': 0,
        'Clustering (Louvain y KMeans)': 0,
        'UserLogic': 0,
        'DataAccessObject': 0,
        'TextAnalysis': 0,
        'WordFrequencyAnalysis': 0,
        'StatisticalAnalysis': 0,
        'FileGeneration': 0,
        'ScheduledTaskManagement': 0
    }
}

for filename in os.listdir(path):
    print(filename)
    fields = filename.split(" - ")
    # print(fields)
    date = fields[0]
    name = fields[1]
    print(fields)
    print(date)
    print(name)
    date_fields = date.split("-")
    print(date_fields)
    date = date_fields[0]
    date_fields_aux = date.split("_")
    year = date_fields_aux[0]
    month = date_fields_aux[1]
    if month == "11":
        month = "noviembre"
    elif month == "12":
        month = "diciembre"
    elif month == "01":
        month = "enero"
    elif month == "02":
        month = "febrero"
    elif month == "03":
        month = "marzo"
    elif month == "04":
        month = "abril"
    else:
        month = "no determinado"
    day = date_fields_aux[2]
    print(name)
    # falta if server o local entonces algo pasa para esos archivos que son resultados ya sea en el servidor o locales
    server = 0
    local = 0
    if "Server" in filename or "server" in filename:
        location = "en el servidor Hércules"
        server += 1
    else:
        location = "de forma local"
        local += 1
    if "Performance" in filename or 'performance' in filename:
        print('execution')
        # It is a execution
        component_name = filename
        if "CL" in component_name:
            component_name = "ConversationLogic"
        elif "Analytics" in component_name:
            component_name = "Clustering (Louvain y KMeans)"
        elif 'UL' in component_name:
            component_name = "UserLogic"
        elif 'delete' in component_name or 'drop' in component_name or 'create_index' in component_name or 'count' in component_name:
            component_name = "DataAccessObject"
        elif "TextAnalysis" in component_name:
            component_name = "TextAnalysis"
        elif "Conversation_Sector_Data" in component_name:
            component_name = "ConversationLogic para verificar funcionamiento de variable conversation_food_time"
        elif "WordFrequencyAnalysis" in component_name:
            component_name = "WordFrequencyAnalysis"
        elif "StatisticalAnalysis" in component_name:
            component_name = "StatisticalAnalysis"
        elif "FileGeneration" in component_name:
            component_name = "FileGeneration"
        elif "ScheduledTaskManagement" in component_name:
            component_name = "ScheduledTaskManagement"
        else:
            component_name = component_name
        if component_name in counter[location].keys():
            counter[location][component_name] += 1
        # print(component_name)
        time = date_fields[1]
        time_fields = time.split("_")
        # print(time_fields)
        hour = time_fields[0]
        minute = time_fields[1]
        second = time_fields[2]
        description = "Resultados de rendimiento de la prueba " + location + " para el componente " \
                      + component_name + ", de forma automática ejecutada el " + day + " de " + month + " de " + year \
                      + " a las " + hour + ":" + minute + ":" + second + "."
    elif "Enviroment" in name:
        print('generación de ambiente')
        if 'UserLogic' in name:
            library = 'UserLogic'
        elif 'Clustering' in name:
            library = 'Clustering'
        elif 'StatisticalAnalysis' in name:
            library = 'StatisticalAnalysis'
        elif 'ScheduledTaskManagement' in name:
            library = 'ScheduledTaskManagement'
        elif 'WordFrequencyAnalysis' in name:
            library = 'WordFrequencyAnalysis'
        else:
            library = "Libreria no reconocida"
        description = "Creación del ambiente " + library + " " + location + " en Anaconda, el " + day + " de " \
                      + month + " de " + year
    elif "users" in name or "conversations" in name:
        print('DAO results')
        if "conversations_per_days" in name:
            result_name = "Número total de conversaciones por día de la semana"
        elif "conversations_per_hour" in name:
            result_name = "Número total de conversaciones por hora del día"
        elif "users" in name:
            result_name = "Total de usuarios por mes"
        else:
            result_name = "Archivo no reconocido"
        # print(component_name)
        time = date_fields[1]
        time_fields = time.split("_")
        # print(time_fields)
        hour = time_fields[0]
        minute = time_fields[1]
        second = time_fields[2]
        description = result_name + " resultado de la ejecución " + location \
                      + " del componente DataAccessObject para la actualización de datos en la presentación a" \
                        " Grupo Nutresa, generado el " + day + " de " + month + " de " + year + " a las " \
                      + hour + ":" + minute + ":" + second + "."
    else:
        print('Docs')
        name_fields = name.split(".")
        doc_type = name_fields[len(name_fields)-1]
        print(doc_type)
        print(name)
        if doc_type == "PNG" or doc_type == "png":
            description = name + ". Versión del " + day + " de " + month + " de " + year + "."
        elif doc_type == "py":
            description = "Desarrollo: " + name + ". Versión del " + day + " de " + month + " de " + year + "."
        else:
            description = "Documento: " + name + ". Versión del " + day + " de " + month + " de " + year + "."
    print(description)
    # p = document.add_paragraph()
    # run = p.add_run(filename + ": ")
    # run.bold = True
    # run.underline = True
    # p.add_run(description)

print(counter)

# document.save('Report - Mar - 2018.docx')




